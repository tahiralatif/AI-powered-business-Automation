from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from docx.shared import Inches, Pt
import os
from agents import function_tool

# ------------------------- Helper: Chart Generator -------------------------
def create_chart(trends, filename="trend_chart.png"):
    import matplotlib.pyplot as plt
    names = [d["name"] for d in trends]
    values = [d["score"] for d in trends]

    plt.figure(figsize=(6, 4))
    plt.bar(names, values, color="#3498db")
    plt.title("Market Analysis Score")
    plt.ylabel("Potential Score")
    plt.xticks(rotation=45, ha="right")
    plt.tight_layout()
    plt.savefig(filename)
    plt.close()
    return filename

# ------------------------- Visualize Data Helper -------------------------
def visualize_data(data: dict, chart_type: str = "bar", title: str = "Data Visualization"):
    """
    Create visualizations from data.
    
    Args:
        data: Dictionary with labels and values
        chart_type: Type of chart (bar, line, pie)
        title: Chart title
    """
    import matplotlib.pyplot as plt
    
    labels = list(data.keys())
    values = list(data.values())
    
    plt.figure(figsize=(10, 6))
    
    if chart_type == "bar":
        plt.bar(labels, values, color="#3498db")
    elif chart_type == "line":
        plt.plot(labels, values, marker='o', color="#2ecc71")
    elif chart_type == "pie":
        plt.pie(values, labels=labels, autopct='%1.1f%%')
    else:
        plt.bar(labels, values, color="#3498db")
    
    plt.title(title)
    plt.xticks(rotation=45, ha="right")
    plt.tight_layout()
    
    filename = f"visualization_{chart_type}.png"
    plt.savefig(filename)
    plt.close()
    
    return filename

# ------------------------- Export Pitch Deck Helper -------------------------
def export_pitchdeck(content: dict, output_name: str = "pitch_deck"):
    """
    Export pitch deck as PDF.
    
    Args:
        content: Dictionary with slide content
        output_name: Output filename
    """
    report_file = f"{output_name}.pdf"
    c = canvas.Canvas(report_file, pagesize=A4)
    width, height = A4
    
    slides = [
        ("Business Name", content.get("business_name", "Startup")),
        ("Problem", content.get("problem", "")),
        ("Solution", content.get("solution", "")),
        ("Market Opportunity", content.get("market", "")),
        ("Revenue Model", content.get("revenue", "")),
        ("Competition", content.get("competition", "")),
        ("Go-To-Market", content.get("gtm", "")),
        ("Team", content.get("team", "")),
        ("Financials", content.get("financials", "")),
        ("Funding Ask", content.get("ask", ""))
    ]
    
    for i, (title, text) in enumerate(slides):
        if i > 0:
            c.showPage()
        
        # Slide header
        c.setFillColor(colors.HexColor("#2c3e50"))
        c.rect(0, height - 60, width, 60, fill=True, stroke=False)
        c.setFillColor(colors.white)
        c.setFont("Helvetica-Bold", 18)
        c.drawString(50, height - 40, title)
        
        # Slide content
        c.setFillColor(colors.black)
        c.setFont("Helvetica", 12)
        y = height - 100
        for line in text.split('\n')[:8]:
            c.drawString(50, y, line[:80])
            y -= 20
    
    c.save()
    return report_file

# ------------------------- Export Trends Report (Production Grade) -------------------------
@function_tool
def export_report(trends: list, format_type: str = "pdf", output_name: str = "business_report"):
    """
    Export professional business report in PDF, Word, or Excel.
    """
    chart_file = create_chart(trends)
    format_type = format_type.lower()

    if format_type == "pdf":
        report_file = f"{output_name}.pdf"
        c = canvas.Canvas(report_file, pagesize=A4)
        width, height = A4

        # Header - Professional Branding
        c.setFillColor(colors.HexColor("#2c3e50"))
        c.rect(0, height - 80, width, 80, fill=True, stroke=False)
        c.setFillColor(colors.white)
        c.setFont("Helvetica-Bold", 24)
        c.drawString(50, height - 50, "AI CO-FOUNDER STRATEGY REPORT")

        # Metadata
        c.setFillColor(colors.black)
        c.setFont("Helvetica", 10)
        c.drawString(50, height - 100, f"Generated on: {os.popen('date /t').read().strip()}")
        c.drawString(50, height - 115, "Subject: Feasibility & Market Analysis")

        # Market Trends Section
        c.setFont("Helvetica-Bold", 16)
        c.setFillColor(colors.HexColor("#3498db"))
        c.drawString(50, height - 160, "1. Market Trends Analysis")

        if os.path.exists(chart_file):
            c.drawImage(chart_file, 50, height - 480, width=500, height=300)

        # Footer
        c.setFont("Helvetica-Oblique", 8)
        c.setFillColor(colors.grey)
        c.drawString(width/2 - 50, 30, "Confidential - Powered by AI Co-founder")

        c.save()
        return report_file

    elif format_type == "word":
        report_file = f"{output_name}.docx"
        doc = Document()

        # Professional Styling
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)

        doc.add_heading("AI Co-founder Strategy Report", 0)
        doc.add_paragraph("Market Insights & Strategic Analysis").italic = True

        doc.add_heading("1. Market Trends", level=1)
        if os.path.exists(chart_file):
            doc.add_picture(chart_file, width=Inches(5))

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Trend Metric'
        hdr_cells[1].text = 'Score'

        for t in trends:
            row_cells = table.add_row().cells
            row_cells[0].text = t["name"]
            row_cells[1].text = str(t["score"])

        doc.save(report_file)
        return report_file

    return f"Unsupported format: {format_type}"
